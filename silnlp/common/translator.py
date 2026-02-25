import logging
from abc import ABC, abstractmethod
from collections import defaultdict
from contextlib import AbstractContextManager
from datetime import date
from itertools import groupby
from pathlib import Path
from typing import DefaultDict, Dict, Generator, Generic, Iterable, List, Optional, Tuple, TypeVar

import docx
import nltk
from iso639 import Lang
from machine.corpora import (
    FileParatextProjectSettingsParser,
    FileParatextProjectTextUpdater,
    ParatextProjectSettings,
    ScriptureRef,
    UpdateUsfmParserHandler,
    UpdateUsfmTextBehavior,
    UsfmFileText,
    UsfmStylesheet,
    parse_usfm,
)
from machine.scripture import VerseRef, is_book_id_valid
from scipy.stats import gmean

from silnlp.common.utils import add_tags_to_sentence
from silnlp.nmt.corpora import CorpusPair

from .corpus import load_corpus, write_corpus
from .paratext import get_book_path, get_iso, get_parent_project_dir, get_project_dir
from .postprocesser import NoDetectedQuoteConventionException, PostprocessHandler, UnknownQuoteConventionException
from .translation_data_structures import DraftGroup, SentenceTranslationGroup, TranslatedDraft
from .usfm_utils import UsfmTextRowCollection

LOGGER = logging.getLogger((__package__ or "") + ".translate")

CONFIDENCE_SUFFIX = ".confidences.tsv"


TVerseKey = TypeVar("TVerseKey")


class ConfidenceFile(ABC, Generic[TVerseKey]):

    def __init__(self, path: Path):
        if not path.name.endswith(CONFIDENCE_SUFFIX):
            raise ValueError(f"Confidence file path must end with {CONFIDENCE_SUFFIX}, got {path.name}")
        self._path = path
        self._trg_draft_file_path = self._get_trg_draft_file_path_from_confidence_path(path)

    @classmethod
    def _get_confidence_file_type(cls, trg_draft_file_path: Path) -> type["ConfidenceFile"]:
        if trg_draft_file_path.name.startswith("test.trg-predictions"):
            return TestConfidenceFile
        ext = trg_draft_file_path.suffix.lower()
        if ext in {".usfm", ".sfm"}:
            return UsfmConfidenceFile
        if ext == ".txt":
            return TxtConfidenceFile
        raise ValueError(
            f"No confidence file type corresponds to trg_draft_file_path {trg_draft_file_path}. "
            f"Expected a trg_draft_file_path starting with 'test.trg-predictions' or ending with .usfm/.sfm/.txt."
        )

    @classmethod
    def from_confidence_file_path(cls, confidence_file_path: Path) -> "ConfidenceFile":
        trg_draft_file_path = cls._get_trg_draft_file_path_from_confidence_path(confidence_file_path)
        file_type = cls._get_confidence_file_type(trg_draft_file_path)
        return file_type(confidence_file_path)

    @classmethod
    def from_draft_file_path(cls, trg_draft_file_path: Path) -> "ConfidenceFile":
        confidence_file_path = trg_draft_file_path.with_suffix(f"{trg_draft_file_path.suffix}{CONFIDENCE_SUFFIX}")
        file_type = cls._get_confidence_file_type(trg_draft_file_path)
        return file_type(confidence_file_path)

    def get_path(self) -> Path:
        return self._path

    def get_verses_path(self) -> Path:
        return self._path.with_suffix(".verses.tsv")

    def get_trg_draft_file_path(self) -> Path:
        return self._trg_draft_file_path

    @staticmethod
    def _get_trg_draft_file_path_from_confidence_path(confidence_file_path: Path) -> Path:
        return confidence_file_path.with_name(confidence_file_path.name.removesuffix(CONFIDENCE_SUFFIX))

    @abstractmethod
    def generate_confidence_files(
        self,
        translated_draft: TranslatedDraft,
        scripture_refs: Optional[List[ScriptureRef]] = None,
    ) -> None:
        pass

    @abstractmethod
    def _parse_verse_key(self, raw_key: str) -> TVerseKey:
        pass

    def verse_confidence_iterator(self) -> Generator[Tuple[TVerseKey, float], None, None]:
        with open(self.get_verses_path(), "r", encoding="utf-8") as f:
            headers = f.readline().strip().split("\t")
            confidence_index = headers.index("Confidence")
            for line in f:
                cols = line.strip().split("\t")
                vref_or_index = cols[0]
                confidence = float(cols[confidence_index])
                yield (self._parse_verse_key(vref_or_index), confidence)

    def get_verse_confidences(self) -> List[Tuple[TVerseKey, float]]:
        return list(self.verse_confidence_iterator())


class UsfmConfidenceFile(ConfidenceFile[VerseRef]):

    def get_chapters_path(self) -> Path:
        return self._path.with_suffix(".chapters.tsv")

    def get_books_path(self) -> Path:
        return self._path.parent / "confidences.books.tsv"

    def _parse_verse_key(self, raw_key: str) -> VerseRef:
        return VerseRef.from_string(raw_key)

    def generate_confidence_files(
        self,
        translated_draft: TranslatedDraft,
        scripture_refs: Optional[List[ScriptureRef]] = None,
    ) -> None:
        if scripture_refs is None:
            raise ValueError("scripture_refs should not be None when generating confidence files for USFM/SFM files.")
        translated_draft.write_confidence_scores_to_file(self._path, scripture_refs)
        translated_draft.write_verse_confidence_scores_to_file(self.get_verses_path(), scripture_refs)
        self.write_chapter_confidence_scores_to_file(translated_draft, scripture_refs)
        self.write_book_confidence_score_to_file(translated_draft, scripture_refs)

    def write_chapter_confidence_scores_to_file(
        self, translated_draft: TranslatedDraft, scripture_refs: List[ScriptureRef]
    ) -> None:
        chapter_confidences: DefaultDict[int, List[float]] = defaultdict(list)
        for vref, confidence in zip(scripture_refs, translated_draft.get_all_sequence_confidence_scores()):
            if not vref.is_verse or confidence is None:
                continue
            chapter_confidences[vref.chapter_num].append(confidence)
        with self.get_chapters_path().open("w", encoding="utf-8", newline="\n") as chapter_confidences_file:
            chapter_confidences_file.write("Chapter\tConfidence\n")
            for chapter, confidences in chapter_confidences.items():
                chapter_confidence = gmean(confidences)
                chapter_confidences_file.write(f"{chapter}\t{chapter_confidence}\n")

    def chapter_confidence_iterator(self) -> Generator[Tuple[int, float], None, None]:
        with open(self.get_chapters_path(), "r", encoding="utf-8") as f:
            headers = f.readline().strip().split("\t")
            confidence_index = headers.index("Confidence")
            for line in f:
                cols = line.strip().split("\t")
                chapter = int(cols[0])
                confidence = float(cols[confidence_index])
                yield (chapter, confidence)

    def get_chapter_confidences(self) -> List[Tuple[int, float]]:
        return list(self.chapter_confidence_iterator())

    def write_book_confidence_score_to_file(
        self, translated_draft: TranslatedDraft, scripture_refs: List[ScriptureRef]
    ) -> None:
        book_confidences: List[float] = []
        for vref, confidence in zip(scripture_refs, translated_draft.get_all_sequence_confidence_scores()):
            if not vref.is_verse or confidence is None:
                continue
            book_confidences.append(confidence)

        existing_books: Dict[str, float] = {}
        if self.get_books_path().exists():
            for book, confidence in self.book_confidence_iterator():
                existing_books[book] = confidence

        current_book = scripture_refs[0].book
        existing_books[current_book] = gmean(book_confidences)
        with self.get_books_path().open("w", encoding="utf-8", newline="\n") as book_confidences_file:
            book_confidences_file.write("Book\tConfidence\n")
            for book, confidence in existing_books.items():
                book_confidences_file.write(f"{book}\t{confidence}\n")

    def book_confidence_iterator(self) -> Generator[Tuple[str, float], None, None]:
        with open(self.get_books_path(), "r", encoding="utf-8") as f:
            headers = f.readline().strip().split("\t")
            confidence_index = headers.index("Confidence")
            for line in f:
                cols = line.strip().split("\t")
                book = cols[0]
                confidence = float(cols[confidence_index])
                yield (book, confidence)

    def get_book_confidences(self) -> List[Tuple[str, float]]:
        return list(self.book_confidence_iterator())


class TxtConfidenceFile(ConfidenceFile[int]):

    def get_files_path(self) -> Path:
        return self._path.parent / "confidences.files.tsv"

    def _parse_verse_key(self, raw_key: str) -> int:
        return int(raw_key)

    def generate_confidence_files(
        self,
        translated_draft: TranslatedDraft,
        scripture_refs: Optional[List[ScriptureRef]] = None,
    ) -> None:
        translated_draft.write_confidence_scores_to_file(self._path)
        translated_draft.write_verse_confidence_scores_to_file(self.get_verses_path())
        self._write_file_confidence_score_to_file(translated_draft)

    def _write_file_confidence_score_to_file(
        self,
        translated_draft: TranslatedDraft,
    ) -> None:
        existing_files: Dict[str, float] = {}
        if self.get_files_path().exists():
            for file_stem, confidence in self.file_confidence_iterator():
                existing_files[file_stem] = confidence

        existing_files[self._trg_draft_file_path.stem] = gmean(
            translated_draft.get_all_sequence_confidence_scores(exclude_none_type=True)
        )
        with self.get_files_path().open("w", encoding="utf-8", newline="\n") as file_confidences_file:
            file_confidences_file.write("File\tConfidence\n")
            for file_stem, confidence in existing_files.items():
                file_confidences_file.write(f"{file_stem}\t{confidence}\n")

    def file_confidence_iterator(self) -> Generator[Tuple[str, float], None, None]:
        with open(self.get_files_path(), "r", encoding="utf-8") as f:
            headers = f.readline().strip().split("\t")
            confidence_index = headers.index("Confidence")
            for line in f:
                cols = line.strip().split("\t")
                file_stem = cols[0]
                confidence = float(cols[confidence_index])
                yield (file_stem, confidence)


class TestConfidenceFile(ConfidenceFile[int]):
    def get_verses_path(self) -> Path:
        # Use the verse-level scores file created by the test script
        return self._path.with_suffix(".scores.tsv")

    def _parse_verse_key(self, raw_key: str) -> int:
        return int(raw_key)

    def generate_confidence_files(
        self,
        translated_draft: TranslatedDraft,
        scripture_refs: Optional[List[ScriptureRef]] = None,
    ) -> None:
        translated_draft.write_confidence_scores_to_file(self._path)


def generate_confidence_files(
    translated_draft: TranslatedDraft,
    trg_draft_file_path: Path,
    scripture_refs: Optional[List[ScriptureRef]] = None,
) -> None:
    if not translated_draft.has_sequence_confidence_scores():
        LOGGER.warning(
            f"{trg_draft_file_path} was not translated with beam search, "
            f"so confidence scores will not be calculated for this file."
        )
        return

    confidence_file = ConfidenceFile.from_draft_file_path(trg_draft_file_path)
    confidence_file.generate_confidence_files(translated_draft, scripture_refs)


class Translator(AbstractContextManager["Translator"], ABC):
    @abstractmethod
    def translate(
        self,
        sentences: Iterable[str],
        src_iso: str,
        trg_iso: str,
        produce_multiple_translations: bool = False,
        vrefs: Optional[Iterable[VerseRef]] = None,
    ) -> Generator[SentenceTranslationGroup, None, None]:
        pass

    def translate_text(
        self,
        src_file_path: Path,
        trg_file_path: Path,
        src_iso: str,
        trg_iso: str,
        produce_multiple_translations: bool = False,
        save_confidences: bool = False,
        trg_prefix: str = "",
        tags: Optional[List[str]] = None,
    ) -> None:

        sentences = [add_tags_to_sentence(tags, sentence) for sentence in load_corpus(src_file_path)]
        sentence_translation_groups: List[SentenceTranslationGroup] = list(
            self.translate(sentences, src_iso, trg_iso, produce_multiple_translations)
        )
        draft_set = DraftGroup(sentence_translation_groups)
        for draft_index, translated_draft in enumerate(draft_set.get_drafts(), 1):
            if produce_multiple_translations:
                trg_draft_file_path = trg_file_path.with_suffix(f".{draft_index}{trg_file_path.suffix}")
            else:
                trg_draft_file_path = trg_file_path
            write_corpus(trg_draft_file_path, translated_draft.get_all_translations())

            if save_confidences:
                generate_confidence_files(
                    translated_draft,
                    trg_draft_file_path,
                )

    def translate_book(
        self,
        src_project: str,
        book: str,
        output_path: Path,
        trg_iso: str,
        produce_multiple_translations: bool = False,
        save_confidences: bool = False,
        chapters: List[int] = [],
        trg_project: Optional[str] = None,
        postprocess_handler: PostprocessHandler = PostprocessHandler(),
        experiment_ckpt_str: str = "",
        training_corpus_pairs: List[CorpusPair] = [],
        tags: Optional[List[str]] = None,
    ) -> None:
        book_path = get_book_path(src_project, book)
        if not book_path.is_file():
            raise RuntimeError(f"Can't find file {book_path} for book {book}")
        else:
            LOGGER.info(f"Found the file {book_path} for book {book}")

        self.translate_usfm(
            book_path,
            output_path,
            get_iso(get_project_dir(src_project)),
            trg_iso,
            produce_multiple_translations,
            save_confidences,
            chapters,
            trg_project,
            postprocess_handler,
            experiment_ckpt_str,
            training_corpus_pairs,
            tags,
        )

    def translate_usfm(
        self,
        src_file_path: Path,
        trg_file_path: Path,
        src_iso: str,
        trg_iso: str,
        produce_multiple_translations: bool = False,
        save_confidences: bool = False,
        chapters: List[int] = [],
        trg_project: Optional[str] = None,
        postprocess_handler: PostprocessHandler = PostprocessHandler(),
        experiment_ckpt_str: str = "",
        training_corpus_pairs: List[CorpusPair] = [],
        tags: Optional[List[str]] = None,
    ) -> None:
        # Create UsfmFileText object for source
        src_from_project = False
        src_settings: Optional[ParatextProjectSettings] = None
        stylesheet = UsfmStylesheet("usfm.sty")
        if str(src_file_path).startswith(str(get_project_dir(""))):
            src_from_project = True
            src_settings = FileParatextProjectSettingsParser(src_file_path.parent).parse()
            stylesheet = src_settings.stylesheet
            book_id = src_settings.get_book_id(src_file_path.name)
            assert book_id is not None

            src_file_text = UsfmFileText(
                src_settings.stylesheet,
                src_settings.encoding,
                book_id,
                src_file_path,
                src_settings.versification,
                include_all_text=True,
                project=src_settings.name,
            )
        else:
            # Guess book ID
            with src_file_path.open(encoding="utf-8-sig") as f:
                book_id = f.read().split()[1].upper()
            if not is_book_id_valid(book_id):
                raise ValueError(f"Book ID not detected: {book_id}")

            src_file_text = UsfmFileText(stylesheet, "utf-8-sig", book_id, src_file_path, include_all_text=True)

        sentences = UsfmTextRowCollection(src_file_text, stylesheet, chapters, tags)

        # sentences = [re.sub(" +", " ", add_tags_to_sentence(tags, s.text.strip())) for s in src_file_text]
        # scripture_refs: List[ScriptureRef] = [s.ref for s in src_file_text]
        # vrefs: List[VerseRef] = [sr.verse_ref for sr in scripture_refs]
        LOGGER.info(f"File {src_file_path} parsed correctly.")

        # Filter sentences
        # for i in reversed(range(len(sentences))):
        #    marker = scripture_refs[i].path[-1].name if len(scripture_refs[i].path) > 0 else ""
        #    if (
        #        (len(chapters) > 0 and scripture_refs[i].chapter_num not in chapters)
        #        or marker in PARAGRAPH_TYPE_EMBEDS
        #        or stylesheet.get_tag(marker).text_type == UsfmTextType.NOTE_TEXT
        #    ):
        #        sentences.pop(i)
        #        scripture_refs.pop(i)
        # empty_sents: List[Tuple[int, ScriptureRef]] = []
        # for i in reversed(range(len(sentences))):
        #    if len(sentences[i].strip()) == 0:
        #        sentences.pop(i)
        #        empty_sents.append((i, scripture_refs.pop(i)))

        sentences_to_translate, vrefs = sentences.get_sentences_and_vrefs_for_translation()

        sentence_translation_groups: List[SentenceTranslationGroup] = list(
            self.translate(
                sentences_to_translate,
                src_iso,
                trg_iso,
                produce_multiple_translations,
                vrefs,
            )
        )
        # num_drafts = len(sentence_translation_groups[0])

        # Add empty sentences back in
        # Prevents pre-existing text from showing up in the sections of translated text
        # for idx, vref in reversed(empty_sents):
        #    sentences.insert(idx, "")
        #    scripture_refs.insert(idx, vref)
        #    sentence_translation_groups.insert(idx, [SentenceTranslation("", [], [], None)] * num_drafts)

        text_behavior = (
            UpdateUsfmTextBehavior.PREFER_NEW if trg_project is not None else UpdateUsfmTextBehavior.STRIP_EXISTING
        )

        translated_text_rows = sentences.to_translated_text_row_collection(sentence_translation_groups)

        # draft_set: DraftGroup = DraftGroup(sentence_translation_groups)
        for draft_index, translated_draft in enumerate(translated_text_rows.get_translated_drafts(), 1):
            # postprocess_handler.construct_rows(scripture_refs, sentences, translated_draft.get_all_translations())
            translated_text_rows.construct_postprocessing_rows_for_draft_index(postprocess_handler, draft_index)

            for config in postprocess_handler.configs:

                # Compile draft remarks
                draft_src_str = f"project {src_file_text.project}" if src_from_project else f"file {src_file_path.name}"
                draft_remark = f"This draft of {scripture_refs[0].book} was machine translated on {date.today()} from {draft_src_str} using model {experiment_ckpt_str}. It should be reviewed and edited carefully."
                postprocess_remark = config.get_postprocess_remark()
                remarks = [draft_remark] + ([postprocess_remark] if postprocess_remark else [])

                # Insert translation into the USFM structure of an existing project
                # If the target project is not the same as the translated file's original project,
                # no verses outside of the ones translated will be overwritten
                if trg_project is not None or src_from_project:
                    project_dir = get_project_dir(trg_project if trg_project is not None else src_file_path.parent.name)
                    parent_settings = None
                    parent_project_dir = get_parent_project_dir(project_dir)
                    if parent_project_dir is not None:
                        parent_settings = FileParatextProjectSettingsParser(parent_project_dir).parse()
                    dest_updater = FileParatextProjectTextUpdater(project_dir, parent_settings)
                    usfm_out = dest_updater.update_usfm(
                        book_id=src_file_text.id,
                        rows=config.rows,
                        text_behavior=text_behavior,
                        paragraph_behavior=config.get_paragraph_behavior(),
                        embed_behavior=config.get_embed_behavior(),
                        style_behavior=config.get_style_behavior(),
                        update_block_handlers=(
                            config.create_place_markers_postprocessor().get_update_block_handlers()
                            if config.is_marker_placement_required()
                            else None
                        ),
                        remarks=remarks,
                        compare_segments=True,
                    )

                    if usfm_out is None:
                        raise FileNotFoundError(
                            f"Book {src_file_text.id} does not exist in target project {trg_project}"
                        )
                else:  # Slightly more manual version for updating an individual file
                    with open(src_file_path, encoding="utf-8-sig") as f:
                        usfm = f.read()
                    handler = UpdateUsfmParserHandler(
                        rows=config.rows,
                        id_text=scripture_refs[0].book,
                        text_behavior=text_behavior,
                        paragraph_behavior=config.get_paragraph_behavior(),
                        embed_behavior=config.get_embed_behavior(),
                        style_behavior=config.get_style_behavior(),
                        update_block_handlers=(
                            config.create_place_markers_postprocessor().get_update_block_handlers()
                            if config.is_marker_placement_required()
                            else None
                        ),
                        remarks=remarks,
                    )
                    parse_usfm(usfm, handler)
                    usfm_out = handler.get_usfm()

                if config.is_quotation_mark_denormalization_required():
                    try:
                        quotation_denormalization_postprocessor = (
                            config.create_denormalize_quotation_marks_postprocessor(training_corpus_pairs)
                        )
                        usfm_out = quotation_denormalization_postprocessor.postprocess_usfm(usfm_out)
                    except (UnknownQuoteConventionException, NoDetectedQuoteConventionException) as e:
                        LOGGER.warning(str(e) + " Skipping quotation mark denormalization.")
                        continue

                # Construct output file name write to file
                trg_draft_file_path = trg_file_path.with_stem(trg_file_path.stem + config.get_postprocess_suffix())
                if produce_multiple_translations:
                    trg_draft_file_path = trg_draft_file_path.with_suffix(f".{draft_index}{trg_file_path.suffix}")

                with trg_draft_file_path.open(
                    "w",
                    encoding=(
                        "utf-8"
                        if src_settings is None
                        or src_settings.encoding == "utf-8-sig"
                        or src_settings.encoding == "utf_8_sig"
                        else src_settings.encoding
                    ),
                ) as f:
                    f.write(usfm_out)

                if save_confidences and config.get_postprocess_suffix() == "":
                    generate_confidence_files(translated_draft, trg_draft_file_path, scripture_refs=scripture_refs)

    def translate_docx(
        self,
        src_file_path: Path,
        trg_file_path: Path,
        src_iso: str,
        trg_iso: str,
        produce_multiple_translations: bool = False,
        tags: Optional[List[str]] = None,
    ) -> None:
        nltk.download("punkt")
        tokenizer: nltk.tokenize.PunktSentenceTokenizer
        try:
            src_lang = Lang(src_iso)
            tokenizer = nltk.data.load(f"tokenizers/punkt/{src_lang.name.lower()}.pickle")
        except Exception:
            tokenizer = nltk.data.load("tokenizers/punkt/english.pickle")

        with src_file_path.open("rb") as file:
            doc = docx.Document(file)

        sentences: List[str] = []
        paras: List[int] = []

        for i, paragraph in enumerate(doc.paragraphs):
            for sentence in tokenizer.tokenize(paragraph.text):
                sentences.append(add_tags_to_sentence(tags, sentence))
                paras.append(i)

        draft_set: DraftGroup = DraftGroup(
            list(self.translate(sentences, src_iso, trg_iso, produce_multiple_translations))
        )

        for draft_index, translated_draft in enumerate(draft_set.get_drafts(), 1):
            for para, group in groupby(zip(translated_draft.get_all_translations(), paras), key=lambda t: t[1]):
                text = " ".join(s[0] for s in group)
                doc.paragraphs[para].text = text

            if produce_multiple_translations:
                trg_draft_file_path = trg_file_path.with_suffix(f".{draft_index}{trg_file_path.suffix}")
            else:
                trg_draft_file_path = trg_file_path

            with trg_draft_file_path.open("wb") as file:
                doc.save(file)

    def __enter__(self) -> "Translator":
        return self

    def __exit__(
        self, exc_type, exc_val, exc_tb  # pyright: ignore[reportMissingParameterType, reportUnknownParameterType]
    ) -> None:
        pass
